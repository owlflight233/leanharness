from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from leanharness.errors import PluginError, PluginManifestError
from leanharness.models import ToolCall
from leanharness.permissions import PermissionMode
from leanharness.plugins.contracts import parse_manifest
from leanharness.plugins.host import MAX_PLUGIN_REQUEST_BYTES, PluginHost
from leanharness.plugins.manager import PluginManager
from leanharness.storage import LocalStore
from leanharness.tools import ToolRegistry

DOCX_PLUGIN = Path(__file__).parents[1] / "plugins" / "leanharness-docx"


def test_plugin_manifest_is_strict_and_versioned() -> None:
    value = json.loads((DOCX_PLUGIN / "leanharness-plugin.json").read_text(encoding="utf-8"))
    manifest = parse_manifest(value)
    assert manifest.id == "leanharness-docx"
    assert manifest.protocol_version == "leanharness.plugin.v1"
    assert manifest.tools[0].mutation is True
    with pytest.raises(PluginManifestError):
        parse_manifest({**value, "unknown": True})


def test_plugin_install_enable_disable_and_remove(tmp_path: Path) -> None:
    with LocalStore(tmp_path / "data") as store:
        manager = PluginManager(store)
        installed = manager.install(DOCX_PLUGIN)
        assert installed.id == "leanharness-docx"
        assert installed.enabled is False
        assert Path(installed.install_path).is_dir()
        assert manager.enable(installed.id).enabled is True
        assert manager.disable(installed.id).enabled is False
        manager.remove(installed.id)
        assert manager.list() == []
        assert not Path(installed.install_path).exists()


def test_duplicate_plugin_install_is_rejected_without_replacing_files(tmp_path: Path) -> None:
    with LocalStore(tmp_path / "data") as store:
        manager = PluginManager(store)
        installed = manager.install(DOCX_PLUGIN)
        with pytest.raises(PluginError, match="already installed"):
            manager.install(DOCX_PLUGIN)
        assert Path(installed.install_path).is_dir()


def test_plugin_host_rejects_oversized_input_before_starting_process(tmp_path: Path) -> None:
    host = PluginHost(DOCX_PLUGIN, ("python", "plugin.py"))
    oversized = {"content": "x" * (MAX_PLUGIN_REQUEST_BYTES + 1)}
    with patch("leanharness.plugins.host.subprocess.Popen") as popen, pytest.raises(
        PluginError, match="request exceeded"
    ):
        host.call("docx_generate", oversized, tmp_path / "artifacts")
    popen.assert_not_called()


def test_plugin_host_rejects_unpaired_unicode_surrogate(tmp_path: Path) -> None:
    host = PluginHost(DOCX_PLUGIN, ("python", "plugin.py"))
    with pytest.raises(PluginError, match="invalid Unicode"):
        host.call("docx_generate", {"title": "bad \ud83d"}, tmp_path / "artifacts")


def test_docx_plugin_is_hidden_in_inspect_and_requires_approve(tmp_path: Path) -> None:
    with LocalStore(tmp_path / "data") as store:
        manager = PluginManager(store)
        manager.enable(manager.install(DOCX_PLUGIN).id)
        tools = manager.runtime_tools(tmp_path, ("leanharness-docx",))
        inspect = ToolRegistry(tmp_path, mode=PermissionMode.INSPECT, additional_tools=tools)
        approve = ToolRegistry(tmp_path, mode=PermissionMode.APPROVE, additional_tools=tools)
        assert "docx_generate" not in {item.name for item in inspect.definitions}
        assert "docx_generate" in {item.name for item in approve.definitions}
        assert approve.approval_required(
            ToolCall("call-1", "docx_generate", {"output_path": "report.docx"})
        )


def test_docx_plugin_generates_a_valid_structured_artifact(tmp_path: Path) -> None:
    with LocalStore(tmp_path / "data") as store:
        manager = PluginManager(store)
        manager.enable(manager.install(DOCX_PLUGIN).id)
        registry = ToolRegistry(
            tmp_path,
            mode=PermissionMode.UNRESTRICTED,
            additional_tools=manager.runtime_tools(tmp_path, ("leanharness-docx",)),
        )
        result = registry.execute(
            ToolCall(
                "call-docx",
                "docx_generate",
                {
                    "output_path": "artifacts/report.docx",
                    "title": "LeanHarness 报告 \ud83d\udcc4",
                    "author": "LeanHarness",
                    "subject": "Plugin verification",
                    "header": "Controlled artifact",
                    "footer": "Local only",
                    "sections": [
                        {
                            "heading": "Overview",
                            "level": 1,
                            "blocks": [
                                {"type": "paragraph", "text": "由独立进程生成。"},
                                {"type": "ordered_list", "items": ["Install", "Generate"]},
                                {"type": "unordered_list", "items": ["Safe", "Auditable"]},
                                {"type": "code", "text": "print('ready')", "language": "python"},
                                {
                                    "type": "table",
                                    "headers": ["Capability", "State"],
                                    "rows": [["DOCX", "Ready"]],
                                },
                            ],
                        }
                    ],
                },
            )
        )
        assert result.ok is True
        target = tmp_path / "artifacts" / "report.docx"
        assert target.is_file()
        document = Document(target)
        texts = [paragraph.text for paragraph in document.paragraphs]
        assert "LeanHarness 报告 📄" in texts
        assert "由独立进程生成。" in texts
        assert "Overview" in texts
        assert "print('ready')" in texts
        assert document.tables[0].cell(1, 1).text == "Ready"
        assert document.sections[0].header.paragraphs[0].text == "Controlled artifact"
        assert document.sections[0].footer.paragraphs[0].text == "Local only"
        assert result.public_metadata["path"] == "artifacts/report.docx"
        assert "print('ready')" not in json.dumps(result.public_metadata)
