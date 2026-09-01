"""Standalone JSONL process for the LeanHarness DOCX plugin."""

from __future__ import annotations

import hashlib
import json
import os
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

PROTOCOL = "leanharness.plugin.v1"
MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_SECTIONS = 64
MAX_BLOCKS = 256
MAX_TEXT_CHARS = 200_000


def main() -> int:
    for raw in sys.stdin:
        try:
            request = json.loads(raw)
            response = handle(request)
        except Exception as exc:
            response = {
                "protocol": PROTOCOL,
                "type": "tool.result",
                "ok": False,
                "error": {"code": "DOCX_GENERATION_FAILED", "message": safe_error(exc)},
            }
        print(json.dumps(response, ensure_ascii=False, separators=(",", ":")), flush=True)
    return 0


def handle(request: object) -> dict[str, object]:
    if not isinstance(request, dict) or request.get("protocol") != PROTOCOL:
        raise ValueError("Protocol mismatch")
    message_type = request.get("type")
    if message_type == "initialize":
        return {"protocol": PROTOCOL, "type": "initialized", "ok": True}
    if message_type == "capabilities":
        return {
            "protocol": PROTOCOL,
            "type": "capabilities.result",
            "ok": True,
            "tools": ["docx_generate"],
        }
    if message_type == "shutdown":
        return {"protocol": PROTOCOL, "type": "shutdown.complete", "ok": True}
    if message_type != "tool.call" or request.get("tool") != "docx_generate":
        raise ValueError("Unknown plugin request")
    arguments = request.get("arguments")
    if not isinstance(arguments, dict):
        raise ValueError("Tool arguments must be an object")
    artifact = generate_docx(arguments)
    return {"protocol": PROTOCOL, "type": "tool.result", "ok": True, "artifact": artifact}


def generate_docx(arguments: dict[str, Any]) -> dict[str, object]:
    title = bounded_text(arguments.get("title"), "title", 500)
    output_path = bounded_text(arguments.get("output_path"), "output_path", 500)
    if not output_path.casefold().endswith(".docx"):
        raise ValueError("output_path must end in .docx")
    sections = arguments.get("sections")
    if not isinstance(sections, list) or not 1 <= len(sections) <= MAX_SECTIONS:
        raise ValueError("sections must contain 1 to 64 items")
    artifact_root = Path(os.environ["LEANHARNESS_PLUGIN_ARTIFACT_DIR"]).resolve(strict=True)
    destination = artifact_root / "document.docx"
    document = Document()
    document.core_properties.title = title
    author = arguments.get("author")
    subject = arguments.get("subject")
    if author is not None:
        document.core_properties.author = bounded_text(author, "author", 200)
    if subject is not None:
        document.core_properties.subject = bounded_text(subject, "subject", 500)
    document.add_heading(title, level=0)
    _ensure_code_style(document)
    _set_headers(document, arguments.get("header"), arguments.get("footer"))
    total_chars = len(title)
    block_count = 0
    for section in sections:
        if not isinstance(section, dict) or set(section) - {"heading", "level", "blocks"}:
            raise ValueError("Section is invalid")
        heading = bounded_text(section.get("heading"), "heading", 500)
        level = section.get("level", 1)
        if isinstance(level, bool) or not isinstance(level, int) or not 1 <= level <= 6:
            raise ValueError("Section level must be between 1 and 6")
        document.add_heading(heading, level=level)
        total_chars += len(heading)
        blocks = section.get("blocks")
        if not isinstance(blocks, list):
            raise ValueError("Section blocks must be an array")
        for block in blocks:
            block_count += 1
            if block_count > MAX_BLOCKS:
                raise ValueError("Document contains too many blocks")
            total_chars += add_block(document, block)
            if total_chars > MAX_TEXT_CHARS:
                raise ValueError("Document text exceeds the limit")
    document.save(destination)
    content = destination.read_bytes()
    return {
        "filename": destination.name,
        "media_type": MEDIA_TYPE,
        "byte_size": len(content),
        "sha256": hashlib.sha256(content).hexdigest(),
        "requested_output": output_path,
    }


def add_block(document: Document, block: object) -> int:
    if not isinstance(block, dict):
        raise ValueError("Document block must be an object")
    block_type = block.get("type")
    if block_type == "paragraph":
        text = bounded_text(block.get("text"), "paragraph text", 20_000, allow_empty=True)
        document.add_paragraph(text)
        return len(text)
    if block_type in {"ordered_list", "unordered_list"}:
        items = block.get("items")
        if not isinstance(items, list) or not 1 <= len(items) <= 100:
            raise ValueError("List items are invalid")
        style = "List Number" if block_type == "ordered_list" else "List Bullet"
        length = 0
        for item in items:
            text = bounded_text(item, "list item", 5_000)
            document.add_paragraph(text, style=style)
            length += len(text)
        return length
    if block_type == "code":
        text = bounded_text(block.get("text"), "code", 40_000, allow_empty=True)
        document.add_paragraph(text, style="LeanHarness Code")
        return len(text)
    if block_type == "table":
        headers = block.get("headers")
        rows = block.get("rows", [])
        if not isinstance(headers, list) or not 1 <= len(headers) <= 20:
            raise ValueError("Table headers are invalid")
        if not isinstance(rows, list) or len(rows) > 100:
            raise ValueError("Table rows are invalid")
        clean_headers = [bounded_text(item, "table header", 2_000) for item in headers]
        clean_rows: list[list[str]] = []
        for row in rows:
            if not isinstance(row, list) or len(row) != len(clean_headers):
                raise ValueError("Table row width does not match headers")
            clean_rows.append([bounded_text(item, "table cell", 5_000, allow_empty=True) for item in row])
        table = document.add_table(rows=1, cols=len(clean_headers))
        table.style = "Table Grid"
        for index, value in enumerate(clean_headers):
            table.rows[0].cells[index].text = value
        for row in clean_rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
        return sum(map(len, clean_headers)) + sum(len(cell) for row in clean_rows for cell in row)
    raise ValueError("Document block type is unsupported")


def _ensure_code_style(document: Document) -> None:
    styles = document.styles
    if "LeanHarness Code" in styles:
        return
    style = styles.add_style("LeanHarness Code", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F5F4")
    style.element.get_or_add_pPr().append(shading)


def _set_headers(document: Document, header: object, footer: object) -> None:
    header_text = bounded_text(header, "header", 1_000, allow_empty=True) if header is not None else ""
    footer_text = bounded_text(footer, "footer", 1_000, allow_empty=True) if footer is not None else ""
    for section in document.sections:
        if header_text:
            section.header.paragraphs[0].text = header_text
        if footer_text:
            section.footer.paragraphs[0].text = footer_text


def bounded_text(value: object, name: str, maximum: int, *, allow_empty: bool = False) -> str:
    if not isinstance(value, str) or len(value) > maximum or (not allow_empty and not value.strip()):
        raise ValueError(f"{name} is invalid")
    return value


def safe_error(exc: Exception) -> str:
    if isinstance(exc, (ValueError, KeyError)):
        return str(exc)[:500]
    return "DOCX generation failed safely"


if __name__ == "__main__":
    raise SystemExit(main())
